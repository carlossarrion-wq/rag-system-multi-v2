import os
import json
import fitz  # PyMuPDF
import pandas as pd
from PIL import Image
import base64
from io import BytesIO
from loguru import logger
from typing import List, Dict, Any, Optional
import hashlib
import xml.etree.ElementTree as ET
from xml.dom import minidom
from docx import Document

class DocumentLoader:
    def __init__(self, connection_manager):
        self.conn_manager = connection_manager
        # AÑADIDO: Soporte para XML
        self.supported_extensions = {'.pdf', '.docx', '.xlsx', '.xls', '.txt', '.png', '.jpg', '.jpeg', '.xml'}

    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a single document and extract its content"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()

            if file_extension not in self.supported_extensions:
                logger.warning(f"Unsupported file type: {file_extension}")
                return None

            # Get file metadata
            file_stats = os.stat(file_path)
            file_hash = self._calculate_file_hash(file_path)

            document = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'file_extension': file_extension,
                'file_size': file_stats.st_size,
                'file_hash': file_hash,
                'content': '',
                'images': [],
                'metadata': {}
            }

            # Extract content based on file type
            if file_extension == '.pdf':
                document = self._load_pdf(file_path, document)
            elif file_extension in ['.xlsx', '.xls']:
                document = self._load_excel(file_path, document)
            elif file_extension == '.txt':
                document = self._load_text(file_path, document)
            elif file_extension in ['.png', '.jpg', '.jpeg']:
                document = self._load_image(file_path, document)
            elif file_extension == '.docx':
                document = self._load_docx(file_path, document)
            elif file_extension == '.xml':  # NUEVO: Soporte XML
                document = self._load_xml(file_path, document)

            logger.info(f"Successfully loaded document: {file_path}")
            return document

        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            return None

    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def _load_xml(self, file_path: str, document: Dict) -> Dict:
        """Load XML document and extract structured content"""
        try:
            # Parse XML
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Extract content in multiple formats
            content_parts = []
            
            # 1. Pretty-printed XML structure
            rough_string = ET.tostring(root, 'unicode')
            reparsed = minidom.parseString(rough_string)
            pretty_xml = reparsed.toprettyxml(indent="  ")
            
            content_parts.append("=== ESTRUCTURA XML ===")
            content_parts.append(pretty_xml)
            
            # 2. Extract all text content
            all_text = self._extract_xml_text(root)
            if all_text.strip():
                content_parts.append("\n=== CONTENIDO TEXTUAL ===")
                content_parts.append(all_text)
            
            # 3. Extract key-value pairs and attributes
            key_values = self._extract_xml_key_values(root)
            if key_values:
                content_parts.append("\n=== DATOS ESTRUCTURADOS ===")
                for key, value in key_values.items():
                    content_parts.append(f"{key}: {value}")
            
            # 4. Create searchable summary
            summary = self._create_xml_summary(root, os.path.basename(file_path))
            content_parts.insert(0, summary)
            
            document['content'] = '\n'.join(content_parts)
            document['metadata'].update({
                'xml_root_tag': root.tag,
                'xml_namespace': root.tag.split('}')[0].strip('{') if '}' in root.tag else None,
                'xml_elements_count': len(list(root.iter())),
                'xml_attributes_count': sum(len(elem.attrib) for elem in root.iter()),
                'has_structured_data': True
            })
            
            return document
            
        except ET.ParseError as e:
            logger.error(f"XML parsing error in {file_path}: {e}")
            # Fallback: treat as text file
            return self._load_text(file_path, document)
        except Exception as e:
            logger.error(f"Error loading XML {file_path}: {e}")
            return document

    def _extract_xml_text(self, element) -> str:
        """Extract all text content from XML element recursively"""
        texts = []
        
        # Get element text
        if element.text and element.text.strip():
            texts.append(element.text.strip())
        
        # Get text from all child elements
        for child in element:
            child_text = self._extract_xml_text(child)
            if child_text:
                texts.append(child_text)
            
            # Get tail text
            if child.tail and child.tail.strip():
                texts.append(child.tail.strip())
        
        return ' '.join(texts)

    def _extract_xml_key_values(self, element, prefix='') -> Dict[str, str]:
        """Extract key-value pairs from XML structure"""
        key_values = {}
        
        # Add attributes
        for attr_name, attr_value in element.attrib.items():
            key = f"{prefix}{element.tag}@{attr_name}" if prefix else f"{element.tag}@{attr_name}"
            key_values[key] = attr_value
        
        # Add element text if it's a leaf node
        if element.text and element.text.strip() and len(list(element)) == 0:
            key = f"{prefix}{element.tag}" if prefix else element.tag
            key_values[key] = element.text.strip()
        
        # Process children
        for child in element:
            child_prefix = f"{prefix}{element.tag}." if prefix else f"{element.tag}."
            child_kvs = self._extract_xml_key_values(child, child_prefix)
            key_values.update(child_kvs)
        
        return key_values

    def _create_xml_summary(self, root, filename) -> str:
        """Create a searchable summary of the XML document"""
        summary_parts = [
            f"=== RESUMEN DEL DOCUMENTO XML: {filename} ===",
            f"Documento XML con elemento raíz: {root.tag}",
        ]
        
        # Add namespace info
        if '}' in root.tag:
            namespace = root.tag.split('}')[0].strip('{')
            summary_parts.append(f"Namespace: {namespace}")
        
        # Count elements by type
        element_counts = {}
        for elem in root.iter():
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            element_counts[tag] = element_counts.get(tag, 0) + 1
        
        if element_counts:
            summary_parts.append("Elementos encontrados:")
            for tag, count in sorted(element_counts.items()):
                summary_parts.append(f"  - {tag}: {count} ocurrencias")
        
        # Add key attributes if any
        all_attrs = set()
        for elem in root.iter():
            all_attrs.update(elem.attrib.keys())
        
        if all_attrs:
            summary_parts.append(f"Atributos utilizados: {', '.join(sorted(all_attrs))}")
        
        # Add content preview
        text_content = self._extract_xml_text(root)
        if text_content:
            preview = text_content[:200] + "..." if len(text_content) > 200 else text_content
            summary_parts.append(f"Vista previa del contenido: {preview}")
        
        return '\n'.join(summary_parts)

    def _load_pdf(self, file_path: str, document: Dict) -> Dict:
        """Load PDF document with text and images"""
        try:
            pdf_document = fitz.open(file_path)
            text_content = []
            images = []

            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)

                # Extract text
                text = page.get_text()
                if text.strip():
                    text_content.append(f"--- Página {page_num + 1} ---\n{text}")

                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_document, xref)

                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("png")
                            img_base64 = base64.b64encode(img_data).decode()

                            images.append({
                                'page': page_num + 1,
                                'index': img_index,
                                'data': img_base64,
                                'format': 'png'
                            })
                        pix = None
                    except Exception as e:
                        logger.warning(f"Error extracting image from page {page_num + 1}: {e}")

            document['content'] = '\n\n'.join(text_content)
            document['images'] = images
            document['metadata']['total_pages'] = len(pdf_document)
            document['metadata']['total_images'] = len(images)

            pdf_document.close()
            return document

        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            return document

    def _load_excel(self, file_path: str, document: Dict) -> Dict:
        """Load Excel document"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content_parts = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # Convert DataFrame to text representation
                sheet_content = f"--- Hoja: {sheet_name} ---\n"
                sheet_content += df.to_string(index=False)
                content_parts.append(sheet_content)

            document['content'] = '\n\n'.join(content_parts)
            document['metadata']['sheets'] = excel_file.sheet_names
            document['metadata']['total_sheets'] = len(excel_file.sheet_names)

            return document

        except Exception as e:
            logger.error(f"Error loading Excel {file_path}: {e}")
            return document

    def _load_text(self, file_path: str, document: Dict) -> Dict:
        """Load text document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                document['content'] = file.read()
            return document
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    document['content'] = file.read()
                return document
            except Exception as e:
                logger.error(f"Error loading text file {file_path}: {e}")
                return document

    def _load_image(self, file_path: str, document: Dict) -> Dict:
        """Load image document with descriptive content"""
        try:
            with open(file_path, 'rb') as file:
                img_data = file.read()
                img_base64 = base64.b64encode(img_data).decode()

                # Get image info
                img = Image.open(file_path)

                # Generate descriptive content for the image
                file_name = os.path.basename(file_path)
                image_format = img.format.upper() if img.format else 'UNKNOWN'
                width, height = img.size

                # Create descriptive content that can be indexed
                descriptive_content = f"""Imagen: {file_name}
Formato: {image_format}
Dimensiones: {width}x{height} píxeles
Tipo de documento: Imagen visual

Esta es una imagen que contiene información visual importante. El archivo se llama "{file_name}" y es de tipo {image_format}.
La imagen tiene unas dimensiones de {width} píxeles de ancho por {height} píxeles de alto.

Contenido visual: Esta imagen puede contener diagramas, gráficos, esquemas, flujos de proceso, interfaces de usuario,
capturas de pantalla, o cualquier otro tipo de información visual relevante para el sistema GADEA.

Para análisis detallado del contenido visual, se requiere procesamiento con modelos de visión artificial."""

                document['content'] = descriptive_content
                document['images'] = [{
                    'data': img_base64,
                    'format': img.format.lower() if img.format else 'unknown',
                    'size': img.size
                }]
                document['metadata']['image_format'] = img.format
                document['metadata']['image_size'] = img.size
                document['metadata']['has_visual_content'] = True

            return document

        except Exception as e:
            logger.error(f"Error loading image {file_path}: {e}")
            return document

    def _load_docx(self, file_path: str, document: Dict) -> Dict:
        """Load DOCX document and extract text content"""
        try:
            # Load the DOCX document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_content = []
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                table_text = f"\n--- Tabla {table_count} ---\n"
                
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    
                    if row_cells:
                        table_text += " | ".join(row_cells) + "\n"
                
                if table_text.strip():
                    text_content.append(table_text)
            
            # Join all content
            full_content = '\n\n'.join(text_content)
            
            # Create descriptive header
            file_name = os.path.basename(file_path)
            header = f"=== DOCUMENTO WORD: {file_name} ===\n"
            
            if full_content.strip():
                document['content'] = header + full_content
            else:
                document['content'] = header + "Documento Word sin contenido textual extraíble."
            
            # Add metadata
            document['metadata'].update({
                'docx_paragraphs': paragraph_count,
                'docx_tables': table_count,
                'docx_processed': True,
                'has_structured_content': table_count > 0
            })
            
            logger.info(f"Successfully extracted DOCX content: {paragraph_count} paragraphs, {table_count} tables from {file_path}")
            return document
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            # Fallback to basic implementation
            document['content'] = f"DOCX Document: {os.path.basename(file_path)}\n[Error extracting content: {str(e)}]"
            document['metadata']['docx_error'] = str(e)
            return document

    def load_documents_from_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """Load all supported documents from a directory"""
        documents = []

        if not os.path.exists(directory_path):
            logger.error(f"Directory does not exist: {directory_path}")
            return documents

        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_extension = os.path.splitext(file)[1].lower()

                if file_extension in self.supported_extensions:
                    document = self.load_document(file_path)
                    if document:
                        documents.append(document)

        logger.info(f"Loaded {len(documents)} documents from {directory_path}")
        return documents
